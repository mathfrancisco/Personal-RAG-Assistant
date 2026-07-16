"""Leitura de arquivos → RawDocument. Um loader por formato (PDF/DOCX/MD/TXT).

`pypdf` e `python-docx` são importados de forma preguiçosa, então ler .md/.txt
não exige essas libs instaladas.
"""

from __future__ import annotations

import hashlib
from pathlib import Path

from rag_assistant.domain.exceptions import DocumentLoadError, UnsupportedFormatError
from rag_assistant.domain.models import RawDocument

SUPPORTED = {".pdf", ".docx", ".md", ".txt"}


def _sha256(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def load_document(path: str | Path) -> list[RawDocument]:
    """Devolve os RawDocuments de um arquivo (PDF = 1 por página; resto = 1)."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext not in SUPPORTED:
        raise UnsupportedFormatError(f"formato não suportado: {ext} ({p.name})")

    data = p.read_bytes()
    doc_hash = _sha256(data)
    source = str(p)

    try:
        if ext in {".md", ".txt"}:
            text = data.decode("utf-8", errors="replace")
            return [RawDocument(text=text, source=source, doc_hash=doc_hash, file_type=ext)]
        if ext == ".pdf":
            return _load_pdf(p, source, doc_hash)
        if ext == ".docx":
            return _load_docx(p, source, doc_hash)
    except (UnsupportedFormatError, DocumentLoadError):
        raise
    except Exception as exc:  # noqa: BLE001 - normaliza qualquer erro de parsing
        raise DocumentLoadError(f"falha ao ler {p.name}: {exc}") from exc

    raise UnsupportedFormatError(ext)  # unreachable, mas mantém o mypy feliz


def _load_pdf(p: Path, source: str, doc_hash: str) -> list[RawDocument]:
    from pypdf import PdfReader

    reader = PdfReader(str(p))
    docs: list[RawDocument] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if not text:
            continue  # página em branco / imagem sem OCR (fora do escopo V1)
        docs.append(
            RawDocument(text=text, source=source, doc_hash=doc_hash, page=i + 1, file_type=".pdf")
        )
    if not docs:
        raise DocumentLoadError(f"nenhum texto extraído de {p.name} (PDF escaneado?)")
    return docs


def _load_docx(p: Path, source: str, doc_hash: str) -> list[RawDocument]:
    import docx

    document = docx.Document(str(p))
    text = "\n".join(par.text for par in document.paragraphs if par.text.strip())
    return [RawDocument(text=text, source=source, doc_hash=doc_hash, file_type=".docx")]
